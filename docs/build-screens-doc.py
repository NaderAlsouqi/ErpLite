"""Build a screenshots .docx from a pages json + screenshots dir.
Usage: python build-screens-doc.py <pages.json> <shots_dir> <out.docx> [rtl]
Each page: section heading, title, route, description, screenshot (or a
'not available for this user' note if the capture was redirected to login)."""
import json, os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(__file__)
pages_path = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, 'pages.json')
shots_dir  = sys.argv[2] if len(sys.argv) > 2 else os.path.join(ROOT, 'screenshots')
out_path   = sys.argv[3] if len(sys.argv) > 3 else os.path.join(ROOT, 'Pages.docx')
RTL        = len(sys.argv) > 4 and sys.argv[4] in ('1', 'rtl', 'ar')

if not os.path.isabs(shots_dir): shots_dir = os.path.join(ROOT, shots_dir) if not os.path.exists(shots_dir) else shots_dir

pages = json.load(open(pages_path, encoding='utf-8'))
res = {}
rp = os.path.join(shots_dir, '_results.json')
if os.path.exists(rp):
    for r in json.load(open(rp, encoding='utf-8')):
        if isinstance(r, dict) and r.get('route'):
            res[r['route']] = r

L = {
  'cover_title': 'مرجع شاشات النظام' if RTL else 'ErpLite — Pages Reference',
  'intro': 'لقطة شاشة ووصف لكل صفحة في النظام، مأخوذة من التطبيق أثناء التشغيل (تعكس بيانات تجريبية).' if RTL
           else 'Screenshot and description of every page in the system. Captured from the running application; screens reflect sample/test data.',
  'total': 'إجمالي الصفحات: ' if RTL else 'Total pages: ',
  'route': 'المسار: ' if RTL else 'Route: ',
  'na': 'هذه الصفحة غير متاحة لصلاحيات هذا المستخدم.' if RTL else 'This page is not available for this user’s permissions.',
  'err': '⚠ سُجِّل خطأ في هذه الصفحة أثناء الالتقاط — للمراجعة والإصلاح لاحقاً.' if RTL
         else '⚠ An error was detected on this page during capture — to review/fix later.',
}


def rtl_para(p):
    if not RTL:
        return
    pPr = p._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi'); pPr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        rPr = r._r.get_or_add_rPr()
        rtl = OxmlElement('w:rtl'); rtl.set(qn('w:val'), '1'); rPr.append(rtl)
        cs = OxmlElement('w:rFonts'); cs.set(qn('w:cs'), 'Tahoma'); rPr.append(cs)


doc = Document()
doc.styles['Normal'].font.name = 'Tahoma' if RTL else 'Segoe UI'
doc.styles['Normal'].font.size = Pt(10.5)

h = doc.add_heading(L['cover_title'], 0); rtl_para(h)
p = doc.add_paragraph(L['intro']); p.runs[0].italic = True; rtl_para(p)
p = doc.add_paragraph(L['total'] + str(len(pages))); rtl_para(p)
doc.add_page_break()

current = None
for pg in pages:
    s = pg.get('section', '')
    if s != current:
        hs = doc.add_heading(s, 1); rtl_para(hs)
        current = s

    ht = doc.add_heading(pg['title'], 2); rtl_para(ht)

    rp2 = doc.add_paragraph()
    rr = rp2.add_run(L['route']); rr.bold = True
    rc = rp2.add_run(pg['route'])
    rc.font.name = 'Consolas'; rc.font.size = Pt(9.5); rc.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    if RTL:
        pPr = rp2._p.get_or_add_pPr(); pPr.append(OxmlElement('w:bidi')); rp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    dp = doc.add_paragraph(pg['desc']); rtl_para(dp)

    info = res.get(pg['route'])
    if info and info.get('hadError'):
        ep = doc.add_paragraph()
        er = ep.add_run(L['err'])
        er.bold = True; er.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        rtl_para(ep)
    redirected = info and isinstance(info.get('finalUrl'), str) and '/auth/login' in info['finalUrl']
    img = os.path.join(shots_dir, info['file']) if info and info.get('file') else None
    if (not redirected) and img and os.path.exists(img):
        try:
            doc.add_picture(img, width=Inches(6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            np = doc.add_paragraph(f'[screenshot error: {e}]'); rtl_para(np)
    else:
        np = doc.add_paragraph(L['na']);
        np.runs[0].italic = True; np.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x00)
        rtl_para(np)

    doc.add_page_break()

doc.save(out_path)
print('wrote', out_path, 'with', len(pages), 'pages (rtl=%s)' % RTL)
