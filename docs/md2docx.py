import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def shade_para(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    pPr.append(sh)


def add_inline(p, text):
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Consolas'; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            p.add_run(part)


def split_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def main(md_path, docx_path):
    lines = open(md_path, encoding='utf-8').read().split('\n')
    doc = Document()
    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(10.5)

    i, n = 0, len(lines)
    first_h1 = True
    while i < n:
        ln = lines[i]

        # fenced code block
        if ln.strip().startswith('```'):
            i += 1; buf = []
            while i < n and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.15)
            r = p.add_run('\n'.join(buf))
            r.font.name = 'Consolas'; r.font.size = Pt(8.5)
            shade_para(p, 'F2F3F5')
            continue

        # GFM table (header row then |---| separator)
        if ln.strip().startswith('|') and i + 1 < n and re.match(r'^\s*\|?\s*:?-{2,}', lines[i + 1]):
            header = split_row(ln); i += 2
            body = []
            while i < n and lines[i].strip().startswith('|'):
                body.append(split_row(lines[i])); i += 1
            cols = len(header)
            t = doc.add_table(rows=1, cols=cols)
            t.style = 'Table Grid'
            for c, h in enumerate(header):
                cell = t.rows[0].cells[c]
                cell.paragraphs[0].text = ''
                add_inline(cell.paragraphs[0], h)
                for rr in cell.paragraphs[0].runs:
                    rr.bold = True
                shade_cell(cell, 'D9E2F3')
            for row in body:
                cells = t.add_row().cells
                for c in range(cols):
                    cells[c].paragraphs[0].text = ''
                    add_inline(cells[c].paragraphs[0], row[c] if c < len(row) else '')
            doc.add_paragraph()
            continue

        # heading
        m = re.match(r'^(#{1,6})\s+(.*)$', ln)
        if m:
            level = len(m.group(1)); txt = m.group(2).strip()
            if level == 1 and first_h1:
                doc.add_heading(txt, 0); first_h1 = False
            else:
                doc.add_heading(txt, min(level, 4))
            i += 1; continue

        # horizontal rule
        if re.match(r'^\s*([-*_])\1{2,}\s*$', ln):
            i += 1; continue

        # blockquote
        if ln.strip().startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, ln.strip()[1:].strip())
            for rr in p.runs:
                rr.italic = True
                rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1; continue

        # bullet / numbered list
        mb = re.match(r'^\s*[-*]\s+(.*)$', ln)
        if mb:
            p = doc.add_paragraph(style='List Bullet'); add_inline(p, mb.group(1)); i += 1; continue
        mn = re.match(r'^\s*\d+\.\s+(.*)$', ln)
        if mn:
            p = doc.add_paragraph(style='List Number'); add_inline(p, mn.group(1)); i += 1; continue

        # blank
        if not ln.strip():
            i += 1; continue

        # paragraph
        p = doc.add_paragraph()
        add_inline(p, ln.strip())
        i += 1

    doc.save(docx_path)
    print('wrote', docx_path)


if __name__ == '__main__':
    main(sys.argv[1], sys.argv[2])
